
from docx import Document
import xlrd
#currently set to replace all instances, need to add counter to replace certain number of instances
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
                    #break 
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

                
                
# *****************for multiple parts
totalParts = 27
# ****************total number of things to change
variables = 21
filePath = "I:\\Quality\\ECO\\ECO by Year\\ECO 2021\\210104 - Parent Part Update to 01085-001 REV 00B\\"

# list of parts to iterate through
workbook = xlrd.open_workbook("I:\\Quality\\ECO\\ECO by Year\\ECO 2021\\210104 - Parent Part Update to 01085-001 REV 00B\ECO210104MS.xlsx")
worksheet = workbook.sheet_by_name('Sheet1')
                              
for i in range(0,totalParts):
    # please format part numbers in column 1 and rev levels in column 2 or udpate the index below
    partNumber = str(worksheet.cell(i+1, 0).value)
    # handles float to string conversion of part numbers
    
    #rev = worksheet.cell(i+1, 1).value
    rev = "00A"
    # *****************new file path and name
    
    newFilename =  filePath + partNumber + " Rev " + rev + " Label Proof.docx"
    filename = filePath + "labelProof.docx"
    doc = Document(filename)
    
    for j in range(0,variables):
        regex = re.compile(worksheet.cell(0,j).value)
        replace = worksheet.cell(i+1, j).value
        # to fix handling of digits only text, for now add space after digits only
        replace = str(replace)
        if replace.isdigit():
            print(replace)
            replace =int(replace)
            print(replace)
            replace = str(replace)
        docx_replace_regex(doc, regex , replace)
    doc.save(newFilename) 
    print(newFilename+" saved!")
print("Complete")
